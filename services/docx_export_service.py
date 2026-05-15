from __future__ import annotations

import base64
import io
from datetime import datetime, timezone
from typing import Any

from docx import Document


class DocxExportService:
    """Professional DOCX exports for document-engine records."""

    def export(self, *, document: dict[str, Any], title: str | None = None) -> dict[str, Any]:
        doc = Document()
        doc.add_paragraph("Confidential care record")
        doc.add_heading(title or str(document.get("title") or "Care document"), 0)
        doc.add_paragraph(f"Generated {datetime.now(timezone.utc).strftime('%d %b %Y %H:%M UTC')}")

        for key, value in (document.get("sections") or {}).items():
            doc.add_heading(str(key).replace("_", " ").title(), level=1)
            doc.add_paragraph(str(value or "No content recorded."))

        self._append_rows(doc, "Chronology and evidence appendix", document.get("links") or [])
        self._append_rows(doc, "Review history", (document.get("review") or {}).get("events") or [])
        self._append_rows(doc, "Signatures", document.get("signatures") or [])

        buffer = io.BytesIO()
        doc.save(buffer)
        payload = buffer.getvalue()
        return {
            "ok": True,
            "profile": "docx",
            "media_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "filename": self._filename(document),
            "byte_length": len(payload),
            "content_base64": base64.b64encode(payload).decode("ascii"),
            "generated_at": datetime.now(timezone.utc).isoformat(),
            "confidentiality_label": "Confidential care record",
        }

    def _append_rows(self, doc: Document, heading: str, rows: list[dict[str, Any]]) -> None:
        doc.add_heading(heading, level=1)
        if not rows:
            doc.add_paragraph("No linked entries recorded.")
            return
        table = doc.add_table(rows=1, cols=3)
        header = table.rows[0].cells
        header[0].text = "Type"
        header[1].text = "Title"
        header[2].text = "Date / role"
        for row in rows[:50]:
            cells = table.add_row().cells
            cells[0].text = str(row.get("link_type") or row.get("type") or "entry")
            cells[1].text = str(row.get("title") or row.get("summary") or row.get("record_id") or row.get("signed_name") or "Source record")
            cells[2].text = str(row.get("date") or row.get("signed_at") or row.get("role") or "")

    def _filename(self, document: dict[str, Any]) -> str:
        title = str(document.get("title") or "care_document").strip().replace(" ", "_")
        return f"{title}.docx"


docx_export_service = DocxExportService()
