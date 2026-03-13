import os
import re
import tempfile
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from reportlab.lib.colors import HexColor
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer


BRAND_NAME = "IndiCare"
TITLE_COLOR = HexColor("#0f172a")
ACCENT_COLOR = HexColor("#2563eb")
TEXT_COLOR = HexColor("#162033")
MUTED_COLOR = HexColor("#64748b")


def safe_filename(value: str, fallback: str = "ai-note") -> str:
    value = (value or "").strip()
    if not value:
        return fallback

    cleaned = re.sub(r"[^a-zA-Z0-9 _-]+", "", value).strip()
    cleaned = re.sub(r"\s+", "-", cleaned)

    return cleaned[:80] or fallback


def normalise_note_lines(note_text: str) -> list[str]:
    return [line.rstrip() for line in (note_text or "").splitlines()]


def create_docx_export(
    title: str,
    note_text: str,
    template_name: str | None = None
) -> str:
    document = Document()

    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    title_para = document.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_run = title_para.add_run(title or "AI Note")
    title_run.bold = True
    title_run.font.size = Pt(20)
    title_run.font.name = "Arial"

    meta_para = document.add_paragraph()
    meta_run = meta_para.add_run(
        f"{BRAND_NAME} export • {datetime.now().strftime('%d %B %Y, %H:%M')}"
    )
    meta_run.italic = True
    meta_run.font.size = Pt(9)
    meta_run.font.name = "Arial"

    if template_name:
        template_para = document.add_paragraph()
        template_run = template_para.add_run(f"Template: {template_name}")
        template_run.font.size = Pt(10)
        template_run.font.name = "Arial"

    document.add_paragraph("")

    for raw_line in normalise_note_lines(note_text):
        line = raw_line.strip()

        if not line:
            document.add_paragraph("")
            continue

        if line.endswith(":") and len(line) < 80:
            para = document.add_paragraph()
            run = para.add_run(line)
            run.bold = True
            run.font.size = Pt(12)
            run.font.name = "Arial"
            continue

        if line.startswith("•") or line.startswith("- "):
            para = document.add_paragraph(style="List Bullet")
            cleaned = line[1:].strip() if line.startswith("•") else line[2:].strip()
            run = para.add_run(cleaned)
            run.font.size = Pt(10.5)
            run.font.name = "Arial"
            continue

        para = document.add_paragraph()
        run = para.add_run(line)
        run.font.size = Pt(10.5)
        run.font.name = "Arial"

    filename = f"{safe_filename(title)}.docx"
    output_path = os.path.join(tempfile.gettempdir(), filename)
    document.save(output_path)
    return output_path


def create_pdf_export(
    title: str,
    note_text: str,
    template_name: str | None = None
) -> str:
    filename = f"{safe_filename(title)}.pdf"
    output_path = os.path.join(tempfile.gettempdir(), filename)

    doc = SimpleDocTemplate(
        output_path,
        pagesize=A4,
        leftMargin=18 * mm,
        rightMargin=18 * mm,
        topMargin=18 * mm,
        bottomMargin=18 * mm,
        title=title or "AI Note"
    )

    styles = getSampleStyleSheet()

    title_style = ParagraphStyle(
        "IndiCareTitle",
        parent=styles["Title"],
        fontName="Helvetica-Bold",
        fontSize=20,
        leading=24,
        textColor=TITLE_COLOR,
        spaceAfter=8,
        alignment=0
    )

    meta_style = ParagraphStyle(
        "IndiCareMeta",
        parent=styles["Normal"],
        fontName="Helvetica-Oblique",
        fontSize=9,
        leading=12,
        textColor=MUTED_COLOR,
        spaceAfter=6
    )

    heading_style = ParagraphStyle(
        "IndiCareHeading",
        parent=styles["Heading2"],
        fontName="Helvetica-Bold",
        fontSize=12,
        leading=15,
        textColor=ACCENT_COLOR,
        spaceBefore=8,
        spaceAfter=5
    )

    body_style = ParagraphStyle(
        "IndiCareBody",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=10.5,
        leading=15,
        textColor=TEXT_COLOR,
        spaceAfter=6
    )

    bullet_style = ParagraphStyle(
        "IndiCareBullet",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=10.5,
        leading=15,
        textColor=TEXT_COLOR,
        leftIndent=12,
        firstLineIndent=-6,
        spaceAfter=4
    )

    story = []

    safe_title = escape_pdf_text(title or "AI Note")
    story.append(Paragraph(safe_title, title_style))
    story.append(
        Paragraph(
            escape_pdf_text(f"{BRAND_NAME} export • {datetime.now().strftime('%d %B %Y, %H:%M')}"),
            meta_style
        )
    )

    if template_name:
        story.append(
            Paragraph(
                escape_pdf_text(f"Template: {template_name}"),
                meta_style
            )
        )

    story.append(Spacer(1, 6))

    for raw_line in normalise_note_lines(note_text):
        line = raw_line.strip()

        if not line:
            story.append(Spacer(1, 4))
            continue

        if line.endswith(":") and len(line) < 80:
            story.append(Paragraph(escape_pdf_text(line), heading_style))
            continue

        if line.startswith("•"):
            story.append(Paragraph(escape_pdf_text(line), bullet_style))
            continue

        if line.startswith("- "):
            story.append(Paragraph(escape_pdf_text(f"• {line[2:].strip()}"), bullet_style))
            continue

        story.append(Paragraph(escape_pdf_text(line), body_style))

    doc.build(story)
    return output_path


def escape_pdf_text(text: str) -> str:
    return (
        str(text or "")
        .replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )
