import io
from html import escape

from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas


def _fmt(value):
    if value is None or value == "":
        return "—"
    return str(value)


def _pairs_from_row(row):
    return [(k, row.get(k)) for k in row.keys() if row.get(k) not in (None, "", [], {})]


def _sections_for_record(record_type: str, row: dict):
    title = row.get("title") or row.get("topic") or row.get("incident_type") or row.get("report_type") or record_type.replace("_", " ").title()
    meta = []

    if row.get("created_at"):
        meta.append(f"Created: {_fmt(row.get('created_at'))}")
    if row.get("updated_at"):
        meta.append(f"Updated: {_fmt(row.get('updated_at'))}")

    if record_type == "plan":
        title = row.get("title") or row.get("plan_type") or "Plan"
        meta = [
            f"Plan Type: {_fmt(row.get('plan_type'))}",
            f"Status: {_fmt(row.get('status'))}",
            f"Workflow: {_fmt(row.get('workflow_status'))}",
            f"Review Date: {_fmt(row.get('review_date'))}",
        ]
    elif record_type == "risk":
        title = row.get("title") or row.get("category") or "Risk Assessment"
        meta = [
            f"Category: {_fmt(row.get('category'))}",
            f"Severity: {_fmt(row.get('severity'))}",
            f"Likelihood: {_fmt(row.get('likelihood'))}",
            f"Workflow: {_fmt(row.get('workflow_status'))}",
        ]
    elif record_type == "daily_note":
        title = f"Daily Note - {_fmt(row.get('shift_type'))}"
        meta = [
            f"Date: {_fmt(row.get('note_date'))}",
            f"Shift: {_fmt(row.get('shift_type'))}",
            f"Workflow: {_fmt(row.get('workflow_status'))}",
        ]
    elif record_type == "incident":
        title = row.get("incident_type") or "Incident"
        meta = [
            f"Date/Time: {_fmt(row.get('incident_datetime'))}",
            f"Severity: {_fmt(row.get('severity'))}",
            f"Workflow: {_fmt(row.get('workflow_status'))}",
            f"Location: {_fmt(row.get('location'))}",
        ]
    elif record_type == "keywork":
        title = row.get("topic") or "Key Work Session"
        meta = [
            f"Session Date: {_fmt(row.get('session_date'))}",
            f"Status: {_fmt(row.get('status'))}",
            f"Workflow: {_fmt(row.get('workflow_status'))}",
        ]
    elif record_type == "handover":
        title = row.get("title") or "Shift Handover"
        meta = [
            f"Handover Date: {_fmt(row.get('handover_date'))}",
            f"Shift: {_fmt(row.get('shift_type'))}",
            f"Status: {_fmt(row.get('status'))}",
        ]
    elif record_type == "report":
        title = row.get("title") or "AI Report"
        meta = [
            f"Report Type: {_fmt(row.get('report_type'))}",
            f"Review Month: {_fmt(row.get('review_month'))}",
            f"Status: {_fmt(row.get('status'))}",
        ]
    elif record_type == "profile":
        title = f"{_fmt(row.get('first_name'))} {_fmt(row.get('last_name'))}".strip() or "Young Person Profile"
        meta = [
            f"Preferred Name: {_fmt(row.get('preferred_name'))}",
            f"Placement Status: {_fmt(row.get('placement_status'))}",
            f"Risk Level: {_fmt(row.get('summary_risk_level'))}",
        ]
    elif record_type == "statutory_document":
        title = row.get("title") or row.get("document_type") or "Statutory Document"
        meta = [
            f"Document Type: {_fmt(row.get('document_type'))}",
            f"Status: {_fmt(row.get('status'))}",
            f"Issue Date: {_fmt(row.get('issue_date'))}",
            f"Review Date: {_fmt(row.get('review_date'))}",
            f"Expiry Date: {_fmt(row.get('expiry_date'))}",
        ]

    sections = _pairs_from_row(row)
    return title, meta, sections


def render_html_document(record_type: str, row: dict):
    title, meta, sections = _sections_for_record(record_type, row)

    meta_html = "".join(f"<div>{escape(m)}</div>" for m in meta if m)
    sections_html = "".join(
        f"""
        <section class="section">
          <h3>{escape(str(k).replace('_', ' ').title())}</h3>
          <div class="value">{escape(_fmt(v))}</div>
        </section>
        """
        for k, v in sections
    )

    return f"""
    <!doctype html>
    <html lang="en">
    <head>
      <meta charset="utf-8" />
      <title>{escape(title)}</title>
      <style>
        body {{
          font-family: Inter, Arial, sans-serif;
          margin: 40px;
          color: #1e252b;
          background: #fff;
        }}
        h1 {{
          margin: 0 0 12px;
          font-size: 28px;
        }}
        .meta {{
          color: #5f6b76;
          margin-bottom: 28px;
          line-height: 1.7;
        }}
        .section {{
          margin-bottom: 20px;
          border-top: 1px solid #d9dfe4;
          padding-top: 14px;
        }}
        .section h3 {{
          margin: 0 0 8px;
          font-size: 15px;
        }}
        .value {{
          white-space: pre-wrap;
          line-height: 1.7;
        }}
        @media print {{
          body {{ margin: 18mm; }}
        }}
      </style>
    </head>
    <body>
      <div style="font-size:12px;font-weight:700;color:#7f1d1d;margin-bottom:12px;">
        Confidential — authorised children's home record
      </div>
      <h1>{escape(title)}</h1>
      <div class="meta">{meta_html}</div>
      {sections_html}
    </body>
    </html>
    """


def build_docx_bytes(record_type: str, row: dict):
    title, meta, sections = _sections_for_record(record_type, row)
    doc = Document()
    doc.add_paragraph("Confidential — authorised children's home record")
    doc.add_heading(title, 0)

    for m in meta:
      if m:
        doc.add_paragraph(m)

    doc.add_paragraph("")

    for k, v in sections:
        doc.add_heading(str(k).replace("_", " ").title(), level=2)
        doc.add_paragraph(_fmt(v))

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def build_pdf_bytes(record_type: str, row: dict):
    title, meta, sections = _sections_for_record(record_type, row)
    buf = io.BytesIO()
    pdf = canvas.Canvas(buf, pagesize=A4)
    width, height = A4
    y = height - 50

    pdf.setFont("Helvetica-Bold", 18)
    pdf.drawString(40, y, "Confidential - authorised children's home record")
    y -= 24
    pdf.drawString(40, y, title)
    y -= 28

    pdf.setFont("Helvetica", 10)
    for m in meta:
        if y < 50:
            pdf.showPage()
            y = height - 50
        pdf.drawString(40, y, m)
        y -= 14

    y -= 10

    for k, v in sections:
        if y < 90:
            pdf.showPage()
            y = height - 50

        pdf.setFont("Helvetica-Bold", 11)
        pdf.drawString(40, y, str(k).replace("_", " ").title())
        y -= 14

        pdf.setFont("Helvetica", 10)
        text = _fmt(v)
        for line in text.splitlines() or ["—"]:
            if y < 50:
                pdf.showPage()
                y = height - 50
            pdf.drawString(50, y, line[:120])
            y -= 12

        y -= 8

    pdf.save()
    buf.seek(0)
    return buf
