from fastapi import APIRouter, Depends, HTTPException, BackgroundTasks, Request
from fastapi.responses import FileResponse
from pydantic import BaseModel, Field
from openai import OpenAI
from docx import Document
from docx.enum.section import WD_ORIENT
from psycopg2.extras import RealDictCursor
from slowapi import Limiter
from slowapi.util import get_remote_address
import uuid
import json
import os
import tempfile
import logging

from auth.session_user import get_current_user
from db.connection import get_db
from services.os_sync_hooks import sync_after_save

router = APIRouter(
    prefix="/documents",
    tags=["Documents"],
)

logger = logging.getLogger(__name__)
limiter = Limiter(key_func=get_remote_address)
client = OpenAI()

ALLOWED_ROLES = {"manager", "admin"}
MAX_DESCRIPTION_LENGTH = 8000
ALLOWED_DOC_TYPES = {
    "incident",
    "risk",
    "daily-log",
    "handover",
    "safeguarding",
    "reflection",
}


class DocumentRequest(BaseModel):
    description: str = Field(..., min_length=10, max_length=MAX_DESCRIPTION_LENGTH)
    young_person_id: int | None = None
    title: str | None = None


def require_document_access(current_user: dict):
    role = (current_user.get("role") or "").strip().lower()
    if role not in ALLOWED_ROLES:
        raise HTTPException(status_code=403, detail="Manager or admin access required")

    home_id = current_user.get("home_id")
    try:
        home_id = int(home_id) if home_id is not None else None
    except (TypeError, ValueError):
        home_id = None

    if role == "manager" and home_id is None:
        raise HTTPException(status_code=403, detail="Manager is not assigned to a home")

    return role, home_id


def create_temp_docx() -> str:
    fd, path = tempfile.mkstemp(suffix=".docx", prefix="indicare_")
    os.close(fd)
    return path


def remove_file_safely(path: str):
    try:
        if path and os.path.exists(path):
            os.remove(path)
    except Exception:
        logger.exception("Failed to remove temporary file: %s", path)


def save_and_return_doc(
    doc: Document,
    download_name: str,
    background_tasks: BackgroundTasks,
):
    path = create_temp_docx()
    doc.save(path)
    background_tasks.add_task(remove_file_safely, path)
    return FileResponse(
        path,
        filename=download_name,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


def write_simple_doc(title: str, text: str) -> Document:
    doc = Document()
    doc.add_heading(title, level=1)

    for line in (text or "").split("\n"):
        clean = line.strip()
        if clean:
            doc.add_paragraph(clean)

    return doc


def audit_document_action(
    conn,
    admin_user_id: int | None,
    action: str,
    target_type: str,
    details: dict,
) -> int | None:
    try:
        with conn.cursor(cursor_factory=RealDictCursor) as cur:
            cur.execute(
                """
                INSERT INTO admin_audit_log (
                    admin_user_id,
                    action,
                    target_type,
                    target_id,
                    details,
                    created_at
                )
                VALUES (%s, %s, %s, NULL, %s::jsonb, NOW())
                RETURNING id, created_at
                """,
                (
                    admin_user_id,
                    action,
                    target_type,
                    json.dumps(details),
                ),
            )
            row = cur.fetchone()
        conn.commit()
        return int(row["id"]) if row and row.get("id") is not None else None
    except Exception:
        conn.rollback()
        logger.exception("Failed to write admin audit log")
        return None


def sync_generated_document(
    *,
    audit_id: int | None,
    payload: DocumentRequest,
    current_user: dict,
    document_type: str,
    title: str,
    download_name: str,
    model_text: str,
) -> None:
    if not audit_id or not payload.young_person_id:
        return
    record = {
        "id": audit_id,
        "young_person_id": payload.young_person_id,
        "home_id": current_user.get("home_id"),
        "title": payload.title or title,
        "summary": f"Generated {title} document: {payload.description[:240]}",
        "narrative": model_text[:4000],
        "description": payload.description,
        "document_type": document_type,
        "file_name": download_name,
        "status": "generated",
        "workflow_status": "generated",
        "created_by": current_user.get("user_id") or current_user.get("id"),
        "created_by_name": current_user.get("name") or current_user.get("email"),
        "recorded_by_name": current_user.get("name") or current_user.get("email"),
        "quality_standards": quality_standards_for_document(document_type),
        "judgement_areas": judgement_areas_for_document(document_type),
        "standards_rationale": f"Generated {document_type} document linked into OS evidence trail",
        "evidence_strength": "medium",
    }
    sync_after_save(source_table="generated_documents", record=record, recorded_by_name=record.get("recorded_by_name"))


def quality_standards_for_document(document_type: str) -> list[str]:
    mapping = {
        "incident": ["protection_of_children"],
        "risk": ["protection_of_children"],
        "daily-log": ["quality_and_purpose_of_care"],
        "handover": ["leadership_and_management"],
        "safeguarding": ["protection_of_children"],
        "reflection": ["leadership_and_management"],
    }
    return mapping.get(document_type, ["quality_and_purpose_of_care"])


def judgement_areas_for_document(document_type: str) -> list[str]:
    if document_type in {"incident", "risk", "safeguarding"}:
        return ["helped_and_protected"]
    if document_type in {"handover", "reflection"}:
        return ["leadership_and_management"]
    return ["experiences_and_progress"]


def safe_model_text(prompt: str) -> str:
    try:
        res = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {
                    "role": "system",
                    "content": (
                        "You write professional, factual UK residential childcare documents. "
                        "Do not include markdown fences. Do not invent facts. "
                        "Use clear headings and plain professional language."
                    ),
                },
                {"role": "user", "content": prompt},
            ],
            temperature=0.2,
        )
        return (res.choices[0].message.content or "").strip()
    except Exception:
        logger.exception("OpenAI text generation failed")
        raise HTTPException(status_code=502, detail="Document generation service failed")


def safe_model_json(prompt: str) -> list[dict]:
    try:
        res = client.chat.completions.create(
            model="gpt-4o-mini",
            response_format={"type": "json_object"},
            messages=[
                {
                    "role": "system",
                    "content": (
                        "Return valid JSON only. "
                        "Output an object with key 'risks' containing a list of exactly 5 risk items."
                    ),
                },
                {"role": "user", "content": prompt},
            ],
            temperature=0.1,
        )
        raw = (res.choices[0].message.content or "").strip()
        parsed = json.loads(raw)

        risks = parsed.get("risks")
        if not isinstance(risks, list):
            raise ValueError("Missing risks list")

        cleaned = []
        for item in risks[:5]:
            if not isinstance(item, dict):
                continue
            cleaned.append(
                {
                    "hazard": str(item.get("hazard", "")),
                    "who": str(item.get("who", "")),
                    "harm": str(item.get("harm", "")),
                    "likelihood": str(item.get("likelihood", "")),
                    "severity": str(item.get("severity", "")),
                    "controls": str(item.get("controls", "")),
                    "further_controls": str(item.get("further_controls", "")),
                }
            )

        if not cleaned:
            raise ValueError("No usable risk rows returned")

        return cleaned

    except Exception:
        logger.exception("OpenAI JSON generation failed")
        raise HTTPException(status_code=502, detail="Risk assessment generation failed")


def build_context_header(current_user: dict) -> str:
    role = (current_user.get("role") or "").strip()
    home_id = current_user.get("home_id")
    return f"Author role: {role}\nHome ID: {home_id}\n"


@router.post("/incident")
@limiter.limit("10/minute")
def generate_incident(
    request: Request,
    payload: DocumentRequest,
    background_tasks: BackgroundTasks,
    conn=Depends(get_db),
    current_user=Depends(get_current_user),
):
    require_document_access(current_user)

    prompt = f"""
{build_context_header(current_user)}
Write a professional incident report for a UK residential children's home.

Sections:
Incident Summary
Antecedents
Behaviour Observed
Staff Response
Safeguarding Considerations
Outcome
Follow Up Actions

Situation:
{payload.description}
"""
    text = safe_model_text(prompt)
    doc = write_simple_doc("Incident Report", text)

    audit_id = audit_document_action(
        conn,
        int(current_user["user_id"]),
        "generate_incident_document",
        "document_generation",
        {"home_id": current_user.get("home_id"), "young_person_id": payload.young_person_id, "document_type": "incident"},
    )
    sync_generated_document(audit_id=audit_id, payload=payload, current_user=current_user, document_type="incident", title="Incident Report", download_name="Incident_Report.docx", model_text=text)

    return save_and_return_doc(doc, "Incident_Report.docx", background_tasks)


@router.post("/risk")
@limiter.limit("10/minute")
def generate_risk(
    request: Request,
    payload: DocumentRequest,
    background_tasks: BackgroundTasks,
    conn=Depends(get_db),
    current_user=Depends(get_current_user),
):
    require_document_access(current_user)

    prompt = f"""
{build_context_header(current_user)}
Create a risk assessment for a residential children's home.

Return JSON as:
{{
  "risks": [
    {{
      "hazard": "",
      "who": "",
      "harm": "",
      "likelihood": "",
      "severity": "",
      "controls": "",
      "further_controls": ""
    }}
  ]
}}

Situation:
{payload.description}
"""
    risks = safe_model_json(prompt)

    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width

    doc.add_heading("Risk Assessment", level=1)
    table = doc.add_table(rows=1, cols=10)

    headers = [
        "Hazard",
        "Who at Risk",
        "Potential Harm",
        "Likelihood",
        "Severity",
        "Risk Score",
        "Existing Controls",
        "Further Controls",
        "Responsible",
        "Review Date",
    ]

    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h

    for r in risks:
        row = table.add_row().cells
        row[0].text = r["hazard"]
        row[1].text = r["who"]
        row[2].text = r["harm"]
        row[3].text = r["likelihood"]
        row[4].text = r["severity"]
        row[5].text = ""
        row[6].text = r["controls"]
        row[7].text = r["further_controls"]
        row[8].text = ""
        row[9].text = ""

    text = "\n".join(f"{item['hazard']}: {item['controls']}" for item in risks)
    audit_id = audit_document_action(
        conn,
        int(current_user["user_id"]),
        "generate_risk_document",
        "document_generation",
        {"home_id": current_user.get("home_id"), "young_person_id": payload.young_person_id, "document_type": "risk"},
    )
    sync_generated_document(audit_id=audit_id, payload=payload, current_user=current_user, document_type="risk", title="Risk Assessment", download_name="Risk_Assessment.docx", model_text=text)

    return save_and_return_doc(doc, "Risk_Assessment.docx", background_tasks)


@router.post("/daily-log")
@limiter.limit("10/minute")
def generate_daily_log(
    request: Request,
    payload: DocumentRequest,
    background_tasks: BackgroundTasks,
    conn=Depends(get_db),
    current_user=Depends(get_current_user),
):
    require_document_access(current_user)

    prompt = f"""
{build_context_header(current_user)}
Write a daily log entry for residential children's home staff.

Sections:
Summary of the Day
Behaviour Observed
Staff Support Provided
Education / Activities
Health and Wellbeing
Any Concerns
Next Actions

Notes:
{payload.description}
"""
    text = safe_model_text(prompt)
    doc = write_simple_doc("Daily Log Entry", text)

    audit_id = audit_document_action(
        conn,
        int(current_user["user_id"]),
        "generate_daily_log_document",
        "document_generation",
        {"home_id": current_user.get("home_id"), "young_person_id": payload.young_person_id, "document_type": "daily-log"},
    )
    sync_generated_document(audit_id=audit_id, payload=payload, current_user=current_user, document_type="daily-log", title="Daily Log Entry", download_name="Daily_Log.docx", model_text=text)

    return save_and_return_doc(doc, "Daily_Log.docx", background_tasks)


@router.post("/handover")
@limiter.limit("10/minute")
def generate_handover(
    request: Request,
    payload: DocumentRequest,
    background_tasks: BackgroundTasks,
    conn=Depends(get_db),
    current_user=Depends(get_current_user),
):
    require_document_access(current_user)

    prompt = f"""
{build_context_header(current_user)}
Write a shift handover for residential children's home staff.

Sections:
Children Present
Key Events
Safeguarding Concerns
Medication Updates
Appointments
Behaviour Notes
Tasks for Next Shift

Notes:
{payload.description}
"""
    text = safe_model_text(prompt)
    doc = write_simple_doc("Shift Handover", text)

    audit_id = audit_document_action(
        conn,
        int(current_user["user_id"]),
        "generate_handover_document",
        "document_generation",
        {"home_id": current_user.get("home_id"), "young_person_id": payload.young_person_id, "document_type": "handover"},
    )
    sync_generated_document(audit_id=audit_id, payload=payload, current_user=current_user, document_type="handover", title="Shift Handover", download_name="Shift_Handover.docx", model_text=text)

    return save_and_return_doc(doc, "Shift_Handover.docx", background_tasks)


@router.post("/safeguarding")
@limiter.limit("5/minute")
def generate_safeguarding(
    request: Request,
    payload: DocumentRequest,
    background_tasks: BackgroundTasks,
    conn=Depends(get_db),
    current_user=Depends(get_current_user),
):
    require_document_access(current_user)

    prompt = f"""
{build_context_header(current_user)}
Write a safeguarding concern record for a children's home.

Sections:
Concern Description
Child Details
Immediate Actions Taken
Who Was Notified
External Agencies Involved
Outcome

Concern:
{payload.description}
"""
    text = safe_model_text(prompt)
    doc = write_simple_doc("Safeguarding Concern Record", text)

    audit_id = audit_document_action(
        conn,
        int(current_user["user_id"]),
        "generate_safeguarding_document",
        "document_generation",
        {"home_id": current_user.get("home_id"), "young_person_id": payload.young_person_id, "document_type": "safeguarding"},
    )
    sync_generated_document(audit_id=audit_id, payload=payload, current_user=current_user, document_type="safeguarding", title="Safeguarding Concern Record", download_name="Safeguarding_Record.docx", model_text=text)

    return save_and_return_doc(doc, "Safeguarding_Record.docx", background_tasks)


@router.post("/reflection")
@limiter.limit("10/minute")
def generate_reflection(
    request: Request,
    payload: DocumentRequest,
    background_tasks: BackgroundTasks,
    conn=Depends(get_db),
    current_user=Depends(get_current_user),
):
    require_document_access(current_user)

    prompt = f"""
{build_context_header(current_user)}
Write a reflective practice record for children's home staff.

Sections:
Situation
Thoughts and Feelings
Analysis
Learning
Future Actions

Reflection:
{payload.description}
"""
    text = safe_model_text(prompt)
    doc = write_simple_doc("Reflective Practice Record", text)

    audit_id = audit_document_action(
        conn,
        int(current_user["user_id"]),
        "generate_reflection_document",
        "document_generation",
        {"home_id": current_user.get("home_id"), "young_person_id": payload.young_person_id, "document_type": "reflection"},
    )
    sync_generated_document(audit_id=audit_id, payload=payload, current_user=current_user, document_type="reflection", title="Reflective Practice Record", download_name="Reflective_Practice.docx", model_text=text)

    return save_and_return_doc(doc, "Reflective_Practice.docx", background_tasks)
