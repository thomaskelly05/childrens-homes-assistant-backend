from fastapi import APIRouter
from pydantic import BaseModel
from openai import OpenAI
from docx import Document
from docx.shared import Inches
from docx.enum.section import WD_ORIENT
from fastapi.responses import FileResponse
import uuid

router = APIRouter(prefix="/risk", tags=["Risk"])

client = OpenAI()


class RiskRequest(BaseModel):
    description: str


@router.post("/generate")
def generate_risk(payload: RiskRequest):

    prompt = f"""
You help staff in residential children's homes write risk assessments.

Return a structured risk assessment for the following situation.

Provide 5 hazards.

Return JSON format:

[
 {{
 "hazard": "",
 "who_at_risk": "",
 "potential_harm": "",
 "likelihood": "",
 "severity": "",
 "controls": "",
 "further_controls": ""
 }}
]

Situation:
{payload.description}
"""

    res = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[{"role": "user", "content": prompt}]
    )

    import json
    risks = json.loads(res.choices[0].message.content)

    doc = Document()

    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width

    doc.add_heading("IndiCare Risk Assessment", level=1)

    table = doc.add_table(rows=1, cols=10)

    headers = [
        "Hazard",
        "Who at Risk",
        "Potential Harm",
        "Likelihood",
        "Severity",
        "Risk Score",
        "Existing Controls",
        "Further Controls",
        "Responsible",
        "Review Date"
    ]

    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h

    for r in risks:
        row = table.add_row().cells

        row[0].text = r["hazard"]
        row[1].text = r["who_at_risk"]
        row[2].text = r["potential_harm"]
        row[3].text = r["likelihood"]
        row[4].text = r["severity"]
        row[5].text = ""
        row[6].text = r["controls"]
        row[7].text = r["further_controls"]
        row[8].text = ""
        row[9].text = ""

    filename = f"risk_{uuid.uuid4()}.docx"

    path = f"/tmp/{filename}"

    doc.save(path)

    return FileResponse(
        path,
        filename="Risk_Assessment.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
