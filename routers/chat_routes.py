import asyncio
import contextlib
import io
import json
import logging
from typing import Any, Literal

from fastapi import APIRouter, Depends, File, Form, HTTPException, Request, UploadFile
from fastapi.responses import StreamingResponse
from pydantic import BaseModel, Field
from psycopg2.extras import RealDictCursor
from slowapi import Limiter
from slowapi.util import get_remote_address

from auth.current_user import get_current_user
from db.connection import get_db, get_db_connection, release_db_connection
from services.ai_service import generate_ai_stream
from services.assistant_orchestrator import build_assistant_prompt

try:
    from docx import Document
except Exception:
    Document = None

try:
    from pypdf import PdfReader
except Exception:
    PdfReader = None

router = APIRouter(prefix="/chat", tags=["Chat"])
compat_router = APIRouter(tags=["Chat Compatibility"])

logger = logging.getLogger(__name__)
limiter = Limiter(key_func=get_remote_address)

MAX_MESSAGE_CHARS = 20000
MAX_DOCUMENT_CHARS = 120000
MAX_HISTORY_MESSAGES = 12
HEARTBEAT_SECONDS = 15
MAX_FILE_SIZE_BYTES = 5 * 1024 * 1024
HEARTBEAT_MARKER = "__heartbeat__"


class AssistantScope(BaseModel):
    scope_type: Literal["global", "young_person"] = "global"
    home_id: int | None = None
    young_person_id: int | None = None
    record_type: str | None = None
    record_id: int | None = None


class AssistantContext(BaseModel):
    current_view: str | None = None
    young_person_name: str | None = None
    placement_status: str | None = None
    summary_risk_level: str | None = None
    composer_record_type: str | None = None
    home_name: str | None = None
    shift_context: str | None = None


class ChatRequest(BaseModel):
    message: str = Field(min_length=1, max_length=MAX_MESSAGE_CHARS)
    conversation_id: int | None = None
    document_text: str | None = None
    document_name: str | None = None
    response_mode: Literal["quick", "balanced", "deep"] = "balanced"
    scope: AssistantScope | None = None
    context: AssistantContext | None = None


class RenameConversation(BaseModel):
    title: str = Field(min_length=1, max_length=120)


class EditMessagePayload(BaseModel):
    message: str = Field(min_length=1, max_length=MAX_MESSAGE_CHARS)
    document_text: str | None = None
    document_name: str | None = None
    response_mode: Literal["quick", "balanced", "deep"] = "balanced"
    scope: AssistantScope | None = None
    context: AssistantContext | None = None


def clip(text: str | None, max_len: int):
    if not text:
        return None
    return text[:max_len]


def ensure_owner(conn, conversation_id: int, user_id: int):
    with conn.cursor() as cur:
        cur.execute(
            "SELECT id FROM conversations WHERE id = %s AND user_id = %s LIMIT 1",
            (conversation_id, user_id),
        )
        if not cur.fetchone():
            raise HTTPException(status_code=404, detail="Conversation not found")


def get_history(conn, conversation_id: int):
    with conn.cursor(cursor_factory=RealDictCursor) as cur:
        cur.execute(
            """
            SELECT role, message
            FROM (
                SELECT role, message, created_at, id
                FROM messages
                WHERE conversation_id = %s
                ORDER BY created_at DESC, id DESC
                LIMIT %s
            ) t
            ORDER BY created_at ASC, id ASC
            """,
            (conversation_id, MAX_HISTORY_MESSAGES),
        )
        return cur.fetchall()


def get_conversation_messages(conn, conversation_id: int):
    with conn.cursor(cursor_factory=RealDictCursor) as cur:
        cur.execute(
            """
            SELECT id, role, message, created_at
            FROM messages
            WHERE conversation_id = %s
            ORDER BY created_at ASC, id ASC
            """,
            (conversation_id,),
        )
        return cur.fetchall()


def get_doc(conn, conversation_id: int):
    with conn.cursor(cursor_factory=RealDictCursor) as cur:
        cur.execute(
            """
            SELECT id, filename, document_text, created_at
            FROM conversation_documents
            WHERE conversation_id = %s
            ORDER BY created_at DESC, id DESC
            LIMIT 1
            """,
            (conversation_id,),
        )
        return cur.fetchone()


def upsert_document(conn, conversation_id: int, filename: str, document_text: str):
    with conn.cursor(cursor_factory=RealDictCursor) as cur:
        cur.execute(
            """
            SELECT id
            FROM conversation_documents
            WHERE conversation_id = %s
            ORDER BY created_at DESC, id DESC
            LIMIT 1
            """,
            (conversation_id,),
        )
        existing = cur.fetchone()

        if existing:
            cur.execute(
                """
                UPDATE conversation_documents
                SET filename = %s, document_text = %s, created_at = NOW()
                WHERE id = %s
                RETURNING id, filename, document_text, created_at
                """,
                (filename, document_text, existing["id"]),
            )
            return cur.fetchone()

        cur.execute(
            """
            INSERT INTO conversation_documents (conversation_id, filename, document_text)
            VALUES (%s, %s, %s)
            RETURNING id, filename, document_text, created_at
            """,
            (conversation_id, filename, document_text),
        )
        return cur.fetchone()


def delete_document(conn, conversation_id: int):
    with conn.cursor() as cur:
        cur.execute(
            "DELETE FROM conversation_documents WHERE conversation_id = %s",
            (conversation_id,),
        )


def generate_title(message: str) -> str:
    cleaned = " ".join((message or "").strip().split())
    return cleaned[:60] or "New chat"


def save_ai_message(conversation_id: int, text: str):
    conn = None
    try:
        conn = get_db_connection()
        with conn.cursor() as cur:
            cur.execute(
                "INSERT INTO messages (conversation_id, role, message) VALUES (%s, 'assistant', %s)",
                (conversation_id, clip(text, 200000)),
            )
        conn.commit()
    except Exception:
        logger.exception("Save AI message failed for conversation_id=%s", conversation_id)
        if conn and not conn.closed:
            conn.rollback()
    finally:
        release_db_connection(conn)


def extract_text_from_txt(file_bytes: bytes) -> str:
    try:
        return file_bytes.decode("utf-8")
    except UnicodeDecodeError:
        return file_bytes.decode("latin-1", errors="ignore")


def extract_text_from_docx(file_bytes: bytes) -> str:
    if Document is None:
        raise HTTPException(status_code=500, detail="DOCX support is not installed on the server")

    document = Document(io.BytesIO(file_bytes))
    parts = []

    for paragraph in document.paragraphs:
        text = (paragraph.text or "").strip()
        if text:
            parts.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = []
            for cell in row.cells:
                value = (cell.text or "").strip()
                if value:
                    cells.append(value)
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts).strip()


def extract_text_from_pdf(file_bytes: bytes) -> str:
    if PdfReader is None:
        raise HTTPException(status_code=500, detail="PDF support is not installed on the server")

    reader = PdfReader(io.BytesIO(file_bytes))
    parts = []

    for page in reader.pages:
        text = (page.extract_text() or "").strip()
        if text:
            parts.append(text)

    return "\n\n".join(parts).strip()


def extract_document_text(filename: str, file_bytes: bytes) -> str:
    lower = (filename or "").lower()

    if lower.endswith(".txt"):
        return extract_text_from_txt(file_bytes)
    if lower.endswith(".docx"):
        return extract_text_from_docx(file_bytes)
    if lower.endswith(".pdf"):
        text = extract_text_from_pdf(file_bytes)
        if not text.strip():
            raise HTTPException(
                status_code=400,
                detail="No readable text found in PDF. The file may be scanned or image-based.",
            )
        return text

    raise HTTPException(
        status_code=400,
        detail="Unsupported file type. Please upload a .txt, .docx, or .pdf file.",
    )


def sse(data: str):
    safe = (data or "").replace("\r\n", "\n").replace("\r", "\n")
    return "".join(f"data: {line}\n" for line in safe.split("\n")) + "\n"


def sse_event(event: str, payload: Any):
    body = json.dumps(payload, ensure_ascii=False)
    return f"event: {event}\ndata: {body}\n\n"


def sse_done():
    return "event: done\ndata: [DONE]\n\n"


def is_heartbeat(item: Any) -> bool:
    return item == HEARTBEAT_MARKER


async def stream_with_heartbeat(generator):
    task = asyncio.create_task(generator.__anext__())

    try:
        while True:
            done, _ = await asyncio.wait(
                {task},
                timeout=HEARTBEAT_SECONDS,
                return_when=asyncio.FIRST_COMPLETED,
            )

            if task in done:
                try:
                    item = task.result()
                except StopAsyncIteration:
                    break

                if item is not None:
                    yield item

                task = asyncio.create_task(generator.__anext__())
            else:
                yield HEARTBEAT_MARKER

    finally:
        if not task.done():
            task.cancel()
            with contextlib.suppress(Exception):
                await task


def _normalise_sources(value: Any) -> list[dict[str, Any]]:
    if not isinstance(value, list):
        return []

    cleaned: list[dict[str, Any]] = []
    seen: set[str] = set()

    for item in value:
        if not isinstance(item, dict):
            continue

        label = str(item.get("label") or "").strip()
        source_type = str(item.get("type") or "").strip()
        document_title = str(item.get("document_title") or "").strip()
        section = str(item.get("section") or "").strip()
        page_number = item.get("page_number")
        excerpt = str(item.get("excerpt") or "").strip()
        url = item.get("url")

        dedupe_key = f"{label}|{source_type}|{document_title}|{section}|{page_number}|{url}"
        if dedupe_key in seen:
            continue
        seen.add(dedupe_key)

        cleaned.append(
            {
                "type": source_type,
                "label": label,
                "document_title": document_title,
                "section": section,
                "page_number": page_number,
                "excerpt": excerpt,
                "url": url,
            }
        )

    return cleaned


def _normalise_runtime(value: Any) -> dict[str, Any]:
    if not isinstance(value, dict):
        return {}

    runtime = {
        "mode": value.get("mode"),
        "task_type": value.get("task_type"),
        "output_type": value.get("output_type"),
        "urgency": value.get("urgency"),
        "safeguarding_level": value.get("safeguarding_level"),
        "user_role_profile": value.get("user_role_profile"),
        "retrieval_level": value.get("retrieval_level"),
        "reflection_level": value.get("reflection_level"),
        "response_stance": value.get("response_stance"),
        "classification_confidence": value.get("classification_confidence"),
        "secondary_intents": value.get("secondary_intents"),
        "suggested_actions": value.get("suggested_actions") or [],
        "regulation_basis": value.get("regulation_basis") or [],
    }

    return {k: v for k, v in runtime.items() if v not in (None, "", [])}


def _normalise_explainability(value: Any) -> dict[str, Any]:
    if not isinstance(value, dict):
        return {}
    return {k: v for k, v in value.items() if v not in (None, "", [])}


def _normalise_scope_dict(scope: dict[str, Any] | None) -> dict[str, Any]:
    scope = scope or {}
    return {
        "scope_type": str(scope.get("scope_type") or "global").strip().lower() or "global",
        "home_id": scope.get("home_id"),
        "young_person_id": scope.get("young_person_id"),
        "record_type": scope.get("record_type"),
        "record_id": scope.get("record_id"),
    }


def _normalise_context_dict(context: dict[str, Any] | None) -> dict[str, Any]:
    context = context or {}
    return {
        "current_view": context.get("current_view"),
        "young_person_name": context.get("young_person_name"),
        "placement_status": context.get("placement_status"),
        "summary_risk_level": context.get("summary_risk_level"),
        "composer_record_type": context.get("composer_record_type"),
        "home_name": context.get("home_name"),
        "shift_context": context.get("shift_context"),
    }


@router.post("/upload")
@limiter.limit("10/minute")
async def upload_chat_document(
    request: Request,
    file: UploadFile = File(...),
    conversation_id: int | None = Form(default=None),
    conn=Depends(get_db),
    current_user=Depends(get_current_user),
):
    contents = await file.read()

    if len(contents) > MAX_FILE_SIZE_BYTES:
        raise HTTPException(status_code=413, detail="File too large. Maximum size is 5MB.")

    filename = file.filename or "document"
    if not contents:
        raise HTTPException(status_code=400, detail="Uploaded file is empty")

    try:
        extracted_text = extract_document_text(filename, contents)
    except HTTPException:
        raise
    except Exception:
        logger.exception("Document extraction failed")
        raise HTTPException(status_code=500, detail="Could not read the uploaded document")

    extracted_text = clip(extracted_text, MAX_DOCUMENT_CHARS)

    if not extracted_text or not extracted_text.strip():
        raise HTTPException(status_code=400, detail="No readable text was found in the uploaded document")

    if conversation_id is not None:
        ensure_owner(conn, conversation_id, current_user["user_id"])
        upsert_document(conn, conversation_id, filename, extracted_text)

    return {
        "ok": True,
        "filename": filename,
        "text": extracted_text,
        "preview": extracted_text[:1200],
    }


@router.get("/conversations")
def list_conversations(conn=Depends(get_db), current_user=Depends(get_current_user)):
    user_id = current_user["user_id"]

    with conn.cursor(cursor_factory=RealDictCursor) as cur:
        cur.execute(
            """
            SELECT id, title, created_at
            FROM conversations
            WHERE user_id = %s
            ORDER BY created_at DESC, id DESC
            """,
            (user_id,),
        )
        rows = cur.fetchall()

    return rows


@router.get("/conversations/{conversation_id}")
def load_conversation(conversation_id: int, conn=Depends(get_db), current_user=Depends(get_current_user)):
    user_id = current_user["user_id"]
    ensure_owner(conn, conversation_id, user_id)

    rows = get_conversation_messages(conn, conversation_id)
    document = get_doc(conn, conversation_id)

    return {
        "messages": rows,
        "document": {
            "filename": document["filename"],
            "text": document["document_text"],
            "created_at": document["created_at"],
        } if document else None,
    }


@router.delete("/conversations/{conversation_id}/document")
def remove_conversation_document(conversation_id: int, conn=Depends(get_db), current_user=Depends(get_current_user)):
    ensure_owner(conn, conversation_id, current_user["user_id"])
    delete_document(conn, conversation_id)
    return {"ok": True, "message": "Document removed"}


@router.post("")
@router.post("/")
@limiter.limit("20/minute")
async def chat(
    request: Request,
    conn=Depends(get_db),
    current_user=Depends(get_current_user),
):
    user_id = current_user["user_id"]

    try:
        raw = await request.json()
        body = ChatRequest(**raw)
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid request")

    message = body.message.strip()
    conversation_id = body.conversation_id
    document_text = clip(body.document_text, MAX_DOCUMENT_CHARS)
    document_name = body.document_name
    scope = _normalise_scope_dict(body.scope.model_dump() if body.scope else None)
    context = _normalise_context_dict(body.context.model_dump() if body.context else None)

    if not message:
        raise HTTPException(status_code=400, detail="Message required")

    try:
        if not conversation_id:
            with conn.cursor(cursor_factory=RealDictCursor) as cur:
                cur.execute(
                    """
                    INSERT INTO conversations (user_id, title)
                    VALUES (%s, %s)
                    RETURNING id
                    """,
                    (user_id, generate_title(message)),
                )
                conversation_id = cur.fetchone()["id"]
        else:
            ensure_owner(conn, conversation_id, user_id)

        if document_text and document_name:
            upsert_document(conn, conversation_id, document_name, document_text)

        with conn.cursor() as cur:
            cur.execute(
                """
                INSERT INTO messages (conversation_id, role, message)
                VALUES (%s, 'user', %s)
                """,
                (conversation_id, message),
            )

        history = get_history(conn, conversation_id)
        doc = get_doc(conn, conversation_id)

        assistant_prompt_bundle = build_assistant_prompt(
            conn,
            user_id=user_id,
            message=message,
            scope=scope,
            history=history,
            context=context,
        )

    except HTTPException:
        raise
    except Exception:
        logger.exception("Failed preparing chat request")
        raise HTTPException(status_code=500, detail="Could not prepare chat request")

    async def stream():
        ai_text = ""
        sources: list[dict[str, Any]] = []
        runtime: dict[str, Any] = {}
        explainability: dict[str, Any] = {}
        assistant_scope_meta: dict[str, Any] = dict(scope)
        assistant_context_meta: dict[str, Any] = dict(assistant_prompt_bundle.get("context") or context)
        suggested_actions: list[str] = []

        try:
            generator = generate_ai_stream(
                message=assistant_prompt_bundle["prompt"],
                session_id=str(conversation_id),
                history=history,
                document_text=doc["document_text"] if doc else None,
                document_name=doc["filename"] if doc else None,
                response_mode=body.response_mode,
                user_context=assistant_prompt_bundle.get("context") or context,
                user_id=user_id,
                conversation_id=conversation_id,
            )

            async for item in stream_with_heartbeat(generator):
                if is_heartbeat(item):
                    yield ": ping\n\n"
                    continue

                if isinstance(item, str):
                    ai_text += item
                    yield sse(item)
                    continue

                if isinstance(item, dict):
                    item_type = item.get("type")

                    if item_type == "progress":
                        content = str(item.get("content") or "").strip()
                        if content:
                            yield sse_event("progress", {"content": content})
                        continue

                    if item_type == "token":
                        token = str(item.get("content") or "")
                        if token:
                            ai_text += token
                            yield sse(token)
                        continue

                    if item_type == "sources":
                        sources = _normalise_sources(item.get("sources"))
                        continue

                    if item_type == "runtime":
                        runtime = _normalise_runtime(item.get("runtime"))
                        continue

                    if item_type == "explainability":
                        explainability = _normalise_explainability(item.get("explainability"))
                        continue

                    if item_type == "meta":
                        sources = _normalise_sources(item.get("sources")) or sources

                        merged_runtime = _normalise_runtime(item.get("runtime"))
                        if merged_runtime:
                            runtime = merged_runtime

                        merged_explainability = _normalise_explainability(item.get("explainability"))
                        if merged_explainability:
                            explainability = merged_explainability

                        if isinstance(item.get("assistant_scope"), dict):
                            assistant_scope_meta = _normalise_scope_dict(item.get("assistant_scope"))

                        if isinstance(item.get("assistant_context"), dict):
                            assistant_context_meta = item.get("assistant_context") or assistant_context_meta

                        if isinstance(item.get("suggested_actions"), list):
                            suggested_actions = [str(x).strip() for x in item.get("suggested_actions") if str(x).strip()]

                        continue

                    logger.debug("Unhandled AI stream item type=%s", item_type)
                    continue

        except Exception:
            logger.exception("AI stream failed for conversation_id=%s", conversation_id)
            fallback = "Sorry, something went wrong."
            ai_text += fallback
            yield sse(fallback)

        finally:
            if ai_text.strip():
                save_ai_message(conversation_id, ai_text)

            final_runtime = dict(runtime)
            prompt_runtime = assistant_prompt_bundle.get("runtime") or {}
            for key, value in prompt_runtime.items():
                if value not in (None, "", []):
                    final_runtime.setdefault(key, value)

            if suggested_actions:
                existing_actions = final_runtime.get("suggested_actions") or []
                if isinstance(existing_actions, list):
                    merged = []
                    seen = set()
                    for action in [*existing_actions, *suggested_actions]:
                        normalised = str(action).strip()
                        if not normalised:
                            continue
                        lowered = normalised.lower()
                        if lowered in seen:
                            continue
                        seen.add(lowered)
                        merged.append(normalised)
                    final_runtime["suggested_actions"] = merged

            yield sse_event(
                "meta",
                {
                    "conversation_id": conversation_id,
                    "sources": sources,
                    "runtime": final_runtime,
                    "explainability": explainability,
                    "assistant_scope": assistant_scope_meta,
                    "assistant_context": assistant_context_meta,
                    "suggested_actions": suggested_actions,
                },
            )

            yield sse_done()

    return StreamingResponse(
        stream(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache, no-transform",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no",
        },
    )


@router.post("/conversations/{conversation_id}/rename")
def rename_conversation(
    conversation_id: int,
    payload: RenameConversation,
    conn=Depends(get_db),
    current_user=Depends(get_current_user),
):
    ensure_owner(conn, conversation_id, current_user["user_id"])

    title = payload.title.strip()
    if not title:
        raise HTTPException(status_code=400, detail="Title is required")

    with conn.cursor() as cur:
        cur.execute(
            "UPDATE conversations SET title = %s WHERE id = %s",
            (title, conversation_id),
        )

    return {"ok": True, "message": "Conversation renamed"}


@router.delete("/conversations/{conversation_id}")
def delete_conversation(conversation_id: int, conn=Depends(get_db), current_user=Depends(get_current_user)):
    ensure_owner(conn, conversation_id, current_user["user_id"])

    with conn.cursor() as cur:
        cur.execute("DELETE FROM messages WHERE conversation_id = %s", (conversation_id,))
        cur.execute("DELETE FROM conversation_documents WHERE conversation_id = %s", (conversation_id,))
        cur.execute("DELETE FROM conversations WHERE id = %s", (conversation_id,))

    return {"ok": True, "message": "Conversation deleted"}


@router.post("/messages/{message_id}/edit")
@limiter.limit("20/minute")
async def edit_message_and_regenerate(
    request: Request,
    message_id: int,
    payload: EditMessagePayload,
    conn=Depends(get_db),
    current_user=Depends(get_current_user),
):
    user_id = current_user["user_id"]
    new_message = payload.message.strip()
    scope = _normalise_scope_dict(payload.scope.model_dump() if payload.scope else None)
    context = _normalise_context_dict(payload.context.model_dump() if payload.context else None)

    if not new_message:
        raise HTTPException(status_code=400, detail="Message is required")

    try:
        with conn.cursor(cursor_factory=RealDictCursor) as cur:
            cur.execute(
                """
                SELECT m.id, m.conversation_id, m.role
                FROM messages m
                JOIN conversations c ON c.id = m.conversation_id
                WHERE m.id = %s AND c.user_id = %s
                LIMIT 1
                """,
                (message_id, user_id),
            )
            row = cur.fetchone()

        if not row:
            raise HTTPException(status_code=404, detail="Message not found")
        if row["role"] != "user":
            raise HTTPException(status_code=400, detail="Only user messages can be edited")

        conversation_id = row["conversation_id"]
        ensure_owner(conn, conversation_id, user_id)

        if payload.document_text and payload.document_name:
            upsert_document(
                conn,
                conversation_id,
                payload.document_name,
                clip(payload.document_text, MAX_DOCUMENT_CHARS),
            )

        with conn.cursor() as cur:
            cur.execute(
                "UPDATE messages SET message = %s WHERE id = %s",
                (new_message, message_id),
            )
            cur.execute(
                "DELETE FROM messages WHERE conversation_id = %s AND id > %s",
                (conversation_id, message_id),
            )

        history = get_history(conn, conversation_id)
        doc = get_doc(conn, conversation_id)

        assistant_prompt_bundle = build_assistant_prompt(
            conn,
            user_id=user_id,
            message=new_message,
            scope=scope,
            history=history,
            context=context,
        )

    except HTTPException:
        raise
    except Exception:
        logger.exception("Failed preparing regenerate request")
        raise HTTPException(status_code=500, detail="Could not prepare message regeneration")

    async def stream():
        ai_text = ""
        sources: list[dict[str, Any]] = []
        runtime: dict[str, Any] = {}
        explainability: dict[str, Any] = {}
        assistant_scope_meta: dict[str, Any] = dict(scope)
        assistant_context_meta: dict[str, Any] = dict(assistant_prompt_bundle.get("context") or context)
        suggested_actions: list[str] = []

        try:
            generator = generate_ai_stream(
                message=assistant_prompt_bundle["prompt"],
                session_id=str(conversation_id),
                history=history,
                document_text=doc["document_text"] if doc else None,
                document_name=doc["filename"] if doc else None,
                response_mode=payload.response_mode,
                user_context=assistant_prompt_bundle.get("context") or context,
                user_id=user_id,
                conversation_id=conversation_id,
            )

            async for item in stream_with_heartbeat(generator):
                if is_heartbeat(item):
                    yield ": ping\n\n"
                    continue

                if isinstance(item, str):
                    ai_text += item
                    yield sse(item)
                    continue

                if isinstance(item, dict):
                    item_type = item.get("type")

                    if item_type == "progress":
                        content = str(item.get("content") or "").strip()
                        if content:
                            yield sse_event("progress", {"content": content})
                        continue

                    if item_type == "token":
                        token = str(item.get("content") or "")
                        if token:
                            ai_text += token
                            yield sse(token)
                        continue

                    if item_type == "sources":
                        sources = _normalise_sources(item.get("sources"))
                        continue

                    if item_type == "runtime":
                        runtime = _normalise_runtime(item.get("runtime"))
                        continue

                    if item_type == "explainability":
                        explainability = _normalise_explainability(item.get("explainability"))
                        continue

                    if item_type == "meta":
                        sources = _normalise_sources(item.get("sources")) or sources

                        merged_runtime = _normalise_runtime(item.get("runtime"))
                        if merged_runtime:
                            runtime = merged_runtime

                        merged_explainability = _normalise_explainability(item.get("explainability"))
                        if merged_explainability:
                            explainability = merged_explainability

                        if isinstance(item.get("assistant_scope"), dict):
                            assistant_scope_meta = _normalise_scope_dict(item.get("assistant_scope"))

                        if isinstance(item.get("assistant_context"), dict):
                            assistant_context_meta = item.get("assistant_context") or assistant_context_meta

                        if isinstance(item.get("suggested_actions"), list):
                            suggested_actions = [str(x).strip() for x in item.get("suggested_actions") if str(x).strip()]

                        continue

                    logger.debug("Unhandled AI regenerate item type=%s", item_type)
                    continue

        except Exception:
            logger.exception("AI regenerate failed for conversation_id=%s", conversation_id)
            fallback = "Sorry, something went wrong while regenerating."
            ai_text += fallback
            yield sse(fallback)

        finally:
            if ai_text.strip():
                save_ai_message(conversation_id, ai_text)

            final_runtime = dict(runtime)
            prompt_runtime = assistant_prompt_bundle.get("runtime") or {}
            for key, value in prompt_runtime.items():
                if value not in (None, "", []):
                    final_runtime.setdefault(key, value)

            if suggested_actions:
                existing_actions = final_runtime.get("suggested_actions") or []
                if isinstance(existing_actions, list):
                    merged = []
                    seen = set()
                    for action in [*existing_actions, *suggested_actions]:
                        normalised = str(action).strip()
                        if not normalised:
                            continue
                        lowered = normalised.lower()
                        if lowered in seen:
                            continue
                        seen.add(lowered)
                        merged.append(normalised)
                    final_runtime["suggested_actions"] = merged

            yield sse_event(
                "meta",
                {
                    "conversation_id": conversation_id,
                    "sources": sources,
                    "runtime": final_runtime,
                    "explainability": explainability,
                    "assistant_scope": assistant_scope_meta,
                    "assistant_context": assistant_context_meta,
                    "suggested_actions": suggested_actions,
                },
            )

            yield sse_done()

    return StreamingResponse(
        stream(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache, no-transform",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no",
        },
    )


@compat_router.get("/conversations")
def compat_list_conversations(conn=Depends(get_db), current_user=Depends(get_current_user)):
    return list_conversations(conn=conn, current_user=current_user)


@compat_router.get("/conversations/{conversation_id}")
def compat_load_conversation(conversation_id: int, conn=Depends(get_db), current_user=Depends(get_current_user)):
    return load_conversation(conversation_id=conversation_id, conn=conn, current_user=current_user)


@compat_router.delete("/conversations/{conversation_id}")
def compat_delete_conversation(conversation_id: int, conn=Depends(get_db), current_user=Depends(get_current_user)):
    return delete_conversation(conversation_id=conversation_id, conn=conn, current_user=current_user)


@compat_router.post("/conversations/{conversation_id}/rename")
def compat_rename_conversation(
    conversation_id: int,
    payload: RenameConversation,
    conn=Depends(get_db),
    current_user=Depends(get_current_user),
):
    return rename_conversation(
        conversation_id=conversation_id,
        payload=payload,
        conn=conn,
        current_user=current_user,
    )


@compat_router.delete("/conversations/{conversation_id}/document")
def compat_remove_document(conversation_id: int, conn=Depends(get_db), current_user=Depends(get_current_user)):
    return remove_conversation_document(
        conversation_id=conversation_id,
        conn=conn,
        current_user=current_user,
    )
